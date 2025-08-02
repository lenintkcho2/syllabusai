import os
import markdown
from pathlib import Path
from typing import Dict, Any
from docx import Document as DocxDocument
from docx.shared import Inches
from weasyprint import HTML, CSS
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from app.models.content import Content

class FileConverter:
    """Conversor de archivos para generar PDFs y DOCX reales"""
    
    @staticmethod
    def markdown_to_html(markdown_content: str, title: str = "") -> str:
        """Convertir markdown a HTML"""
        
        # CSS mejorado para PDF
        css_styles = """
        <style>
            body {
                font-family: 'Arial', sans-serif;
                line-height: 1.6;
                margin: 2cm;
                color: #333;
            }
            h1 {
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
                margin-bottom: 20px;
            }
            h2 {
                color: #34495e;
                margin-top: 30px;
                margin-bottom: 15px;
            }
            h3 {
                color: #7f8c8d;
                margin-top: 20px;
                margin-bottom: 10px;
            }
            p {
                margin-bottom: 12px;
                text-align: justify;
            }
            ul, ol {
                margin-bottom: 15px;
                padding-left: 25px;
            }
            li {
                margin-bottom: 5px;
            }
            .header {
                text-align: center;
                margin-bottom: 40px;
                border-bottom: 1px solid #ddd;
                padding-bottom: 20px;
            }
            .footer {
                margin-top: 40px;
                text-align: center;
                font-size: 0.9em;
                color: #7f8c8d;
                border-top: 1px solid #ddd;
                padding-top: 20px;
            }
            .section {
                margin-bottom: 25px;
            }
        </style>
        """
        
        # Convertir markdown a HTML
        html_content = markdown.markdown(
            markdown_content,
            extensions=['tables', 'toc', 'codehilite']
        )
        
        # Template HTML completo
        full_html = f"""
        <!DOCTYPE html>
        <html lang="es">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{title}</title>
            {css_styles}
        </head>
        <body>
            <div class="header">
                <h1>{title}</h1>
                <p><em>Generado por SyllabusAI</em></p>
            </div>
            
            <div class="content">
                {html_content}
            </div>
            
            <div class="footer">
                <p>Documento generado automáticamente por SyllabusAI</p>
                <p>Fecha de generación: {Path().absolute()}</p>
            </div>
        </body>
        </html>
        """
        
        return full_html
    
    @staticmethod
    def html_to_pdf(html_content: str, output_path: str, settings: Dict[str, Any]) -> bool:
        """Convertir HTML a PDF usando WeasyPrint"""
        try:
            # Configurar CSS adicional para PDF
            css = CSS(string=f"""
                @page {{
                    size: {settings.get('paper_size', 'A4')};
                    margin: 2cm;
                }}
                body {{
                    font-size: {settings.get('font_size', '12pt')};
                }}
            """)
            
            # Generar PDF
            HTML(string=html_content).write_pdf(
                output_path,
                stylesheets=[css]
            )
            
            return True
            
        except Exception as e:
            print(f"Error generando PDF: {e}")
            return False
    
    @staticmethod
    def markdown_to_docx(markdown_content: str, title: str, output_path: str, settings: Dict[str, Any]) -> bool:
        """Convertir markdown a DOCX usando python-docx"""
        try:
            # Crear documento Word
            doc = DocxDocument()
            
            # Configurar estilos
            styles = doc.styles
            
            # Estilo para título
            title_style = styles.add_style('Custom Title', 1)
            title_style.font.size = Inches(0.2)
            title_style.font.bold = True
            
            # Agregar título
            title_para = doc.add_heading(title, level=1)
            
            # Agregar subtítulo
            subtitle = doc.add_paragraph("Generado por SyllabusAI")
            subtitle.italic = True
            
            # Agregar espacio
            doc.add_paragraph("")
            
            # Procesar contenido markdown línea por línea
            lines = markdown_content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    # Línea vacía - agregar espacio
                    if current_paragraph:
                        current_paragraph = None
                    continue
                
                if line.startswith('# '):
                    # Heading 1
                    doc.add_heading(line[2:], level=1)
                    current_paragraph = None
                elif line.startswith('## '):
                    # Heading 2
                    doc.add_heading(line[3:], level=2)
                    current_paragraph = None
                elif line.startswith('### '):
                    # Heading 3
                    doc.add_heading(line[4:], level=3)
                    current_paragraph = None
                elif line.startswith('- ') or line.startswith('* '):
                    # Lista
                    doc.add_paragraph(line[2:], style='List Bullet')
                    current_paragraph = None
                elif line.startswith('1. ') or line.startswith('2. '):
                    # Lista numerada
                    doc.add_paragraph(line[3:], style='List Number')
                    current_paragraph = None
                else:
                    # Párrafo normal
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph(line)
                    else:
                        current_paragraph.add_run(' ' + line)
            
            # Agregar footer
            doc.add_paragraph("")
            footer = doc.add_paragraph("Documento generado automáticamente por SyllabusAI")
            footer.italic = True
            
            # Guardar documento
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error generando DOCX: {e}")
            return False
    
    @staticmethod
    def generate_pdf(content: Content, output_path: str, settings: Dict[str, Any]) -> bool:
        """Generar PDF real desde contenido"""
        try:
            # Convertir markdown a HTML
            html_content = FileConverter.markdown_to_html(
                content.markdown_content, 
                content.title
            )
            
            # Convertir HTML a PDF
            return FileConverter.html_to_pdf(html_content, output_path, settings)
            
        except Exception as e:
            print(f"Error en generate_pdf: {e}")
            return False
    
    @staticmethod
    def generate_docx(content: Content, output_path: str, settings: Dict[str, Any]) -> bool:
        """Generar DOCX real desde contenido"""
        try:
            return FileConverter.markdown_to_docx(
                content.markdown_content,
                content.title,
                output_path,
                settings
            )
            
        except Exception as e:
            print(f"Error en generate_docx: {e}")
            return False